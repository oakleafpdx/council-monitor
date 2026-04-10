#!/usr/bin/env python3
"""
Council Meeting Pipeline — GitHub Actions Side
===============================================
Receives transcripts from the local script (via repository_dispatch) or
falls back to YouTube captions. Runs Claude summarization and uploads
to Google Drive.
 
Config is externalized:
  - config.json: watch topics, video keywords, channel ID
  - prompt_template.md: Claude prompt (edit to refine summaries)
"""
 
import io
import os
import sys
import json
import argparse
import re
import subprocess
from datetime import datetime, timezone
from pathlib import Path
 
CLAUDE_MODEL = "claude-sonnet-4-20250514"
MAX_TRANSCRIPT_CHARS = 120_000
SCRIPT_DIR = Path(__file__).parent
LEDGER_PATH = SCRIPT_DIR / "processed_videos.json"
 
 
# ---------------------------------------------------------------------------
# Config Loading
# ---------------------------------------------------------------------------
 
def load_config() -> dict:
    config_path = SCRIPT_DIR / "config.json"
    with open(config_path) as f:
        return json.load(f)
 
 
def load_prompt_template() -> str:
    prompt_path = SCRIPT_DIR / "prompt_template.md"
    with open(prompt_path) as f:
        return f.read()
 
 
def load_feedback_as_prompt() -> str:
    """Load feedback.json and format it as prompt instructions."""
    feedback_path = SCRIPT_DIR / "feedback.json"
    if not feedback_path.exists():
        return ""
 
    with open(feedback_path) as f:
        fb = json.load(f)
 
    sections = []
 
    names = fb.get("name_corrections", {})
    if names:
        lines = [f"  - Use \"{correct}\" not \"{wrong}\"" for wrong, correct in names.items()]
        sections.append("NAME CORRECTIONS (always apply these):\n" + "\n".join(lines))
 
    fmt = fb.get("formatting_preferences", [])
    if fmt:
        lines = [f"  - {p}" for p in fmt]
        sections.append("FORMATTING PREFERENCES:\n" + "\n".join(lines))
 
    emphasis = fb.get("topic_emphasis", [])
    if emphasis:
        lines = [f"  - {t}" for t in emphasis]
        sections.append("TOPICS TO EMPHASIZE (provide extra detail on these):\n" + "\n".join(lines))
 
    deemphasis = fb.get("topic_deemphasis", [])
    if deemphasis:
        lines = [f"  - {t}" for t in deemphasis]
        sections.append("TOPICS TO DE-EMPHASIZE (mention briefly or skip):\n" + "\n".join(lines))
 
    instructions = fb.get("standing_instructions", [])
    if instructions:
        lines = [f"  - {i}" for i in instructions]
        sections.append("STANDING INSTRUCTIONS:\n" + "\n".join(lines))
 
    if not sections:
        return ""
 
    return "ANALYST PREFERENCES AND CORRECTIONS:\n" + "\n\n".join(sections) + "\n"
 
 
# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
 
def extract_video_id(url: str) -> str:
    patterns = [
        r'(?:v=|/v/|youtu\.be/)([a-zA-Z0-9_-]{11})',
        r'(?:embed/)([a-zA-Z0-9_-]{11})',
        r'(?:live/)([a-zA-Z0-9_-]{11})',
    ]
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            return match.group(1)
    raise ValueError(f"Could not extract video ID from URL: {url}")
 
 
def format_timestamp(ms: int) -> str:
    total_seconds = ms // 1000
    hours = total_seconds // 3600
    minutes = (total_seconds % 3600) // 60
    seconds = total_seconds % 60
    return f"{hours:02d}:{minutes:02d}:{seconds:02d}"
 
 
# ---------------------------------------------------------------------------
# Git helpers (commit summaries back to repo)
# ---------------------------------------------------------------------------
 
def commit_summary_to_repo(file_paths: list[Path], metadata: dict):
    """Commit summary files (docx + markdown) back to the repo so they persist."""
    try:
        # Configure git identity for the Actions bot
        subprocess.run(
            ["git", "config", "user.name", "github-actions[bot]"],
            check=True, capture_output=True,
        )
        subprocess.run(
            ["git", "config", "user.email", "github-actions[bot]@users.noreply.github.com"],
            check=True, capture_output=True,
        )
 
        # Stage all summary files
        for fp in file_paths:
            subprocess.run(
                ["git", "add", str(fp)],
                check=True, capture_output=True,
            )
 
        # Also stage the updated ledger if it exists
        ledger = Path("processed_videos.json")
        if ledger.exists():
            subprocess.run(
                ["git", "add", str(ledger)],
                check=True, capture_output=True,
            )
 
        # Check if there's anything to commit
        result = subprocess.run(
            ["git", "diff", "--cached", "--quiet"],
            capture_output=True,
        )
        if result.returncode == 0:
            print("[INFO] No changes to commit.")
            return
 
        title = metadata.get("title", "Unknown")
        commit_msg = f"Add summary: {title}"
        subprocess.run(
            ["git", "commit", "-m", commit_msg],
            check=True, capture_output=True,
        )
        subprocess.run(
            ["git", "push"],
            check=True, capture_output=True,
        )
        for fp in file_paths:
            print(f"[INFO] Committed to repo: {fp}")
    except subprocess.CalledProcessError as e:
        stderr = e.stderr.decode() if e.stderr else ""
        print(f"[WARN] Failed to commit summary to repo: {stderr}")
 
 
# ---------------------------------------------------------------------------
# YouTube Data API
# ---------------------------------------------------------------------------
 
def get_youtube_service():
    from googleapiclient.discovery import build
    api_key = os.environ.get("YOUTUBE_API_KEY")
    if not api_key:
        raise ValueError("YOUTUBE_API_KEY not set.")
    return build("youtube", "v3", developerKey=api_key)
 
 
def get_video_metadata(video_id: str) -> dict:
    service = get_youtube_service()
    response = service.videos().list(
        part="snippet,contentDetails", id=video_id
    ).execute()
 
    if not response.get("items"):
        return {"id": video_id, "title": "Unknown", "upload_date": "Unknown", "duration": 0}
 
    item = response["items"][0]
    snippet = item["snippet"]
    duration_str = item["contentDetails"]["duration"]
    match = re.match(r'PT(?:(\d+)H)?(?:(\d+)M)?(?:(\d+)S)?', duration_str)
    duration = 0
    if match:
        duration = (int(match.group(1) or 0) * 3600 +
                    int(match.group(2) or 0) * 60 +
                    int(match.group(3) or 0))
 
    return {
        "id": video_id,
        "title": snippet["title"],
        "upload_date": snippet["publishedAt"][:10],
        "duration": duration,
    }
 
 
def fetch_youtube_captions(video_id: str) -> str:
    """Fallback: fetch YouTube auto-captions."""
    from youtube_transcript_api import YouTubeTranscriptApi
 
    print(f"[INFO] Fetching YouTube captions for: {video_id}")
    transcript_list = YouTubeTranscriptApi.get_transcript(video_id, languages=["en"])
 
    lines = []
    for entry in transcript_list:
        ts = format_timestamp(int(entry["start"] * 1000))
        text = entry["text"].replace("\n", " ")
        lines.append(f"[{ts}] {text}")
    return "\n".join(lines)
 
 
# ---------------------------------------------------------------------------
# Claude Summarization
# ---------------------------------------------------------------------------
 
def build_prompt(
    transcript_text: str,
    video_metadata: dict,
    watch_topics: list[str],
    chapters_text: str = "",
    assemblyai_summary: str = "",
) -> str:
    """Build prompt from template and data."""
    template = load_prompt_template()
    topics_list = "\n".join(f"  - {t}" for t in watch_topics)
 
    assemblyai_section = ""
    if assemblyai_summary:
        assemblyai_section = f"ASSEMBLYAI AUTO-SUMMARY:\n{assemblyai_summary}\n"
 
    chapters_section = ""
    if chapters_text:
        chapters_section = f"AUTO-GENERATED CHAPTERS:\n{chapters_text}\n"
 
    feedback_section = load_feedback_as_prompt()
 
    return template.format(
        title=video_metadata.get("title", "Portland City Council Meeting"),
        upload_date=video_metadata.get("upload_date", "Unknown"),
        duration_min=video_metadata.get("duration", 0) // 60,
        assemblyai_section=assemblyai_section,
        chapters_section=chapters_section,
        feedback_section=feedback_section,
        transcript=transcript_text[:MAX_TRANSCRIPT_CHARS],
        topics_list=topics_list,
    )
 
 
def summarize_with_claude(prompt: str) -> str:
    import anthropic
 
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        raise ValueError("ANTHROPIC_API_KEY not set.")
 
    client = anthropic.Anthropic(api_key=api_key)
    print("[INFO] Sending transcript to Claude for summarization...")
 
    message = client.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=4096,
        messages=[{"role": "user", "content": prompt}],
    )
    return message.content[0].text
 
 
# ---------------------------------------------------------------------------
# Google Drive
# ---------------------------------------------------------------------------
 
def get_drive_service():
    from google.oauth2 import service_account
    from googleapiclient.discovery import build
 
    sa_json = os.environ.get("GOOGLE_SERVICE_ACCOUNT_JSON")
    if not sa_json:
        return None
 
    sa_info = json.loads(sa_json)
    credentials = service_account.Credentials.from_service_account_info(
        sa_info, scopes=["https://www.googleapis.com/auth/drive.file"]
    )
    return build("drive", "v3", credentials=credentials)
 
 
def upload_to_drive(service, filename: str, file_bytes: bytes, folder_id: str) -> str:
    from googleapiclient.http import MediaInMemoryUpload
 
    mimetype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    media = MediaInMemoryUpload(file_bytes, mimetype=mimetype, resumable=True)
    file = service.files().create(
        body={"name": filename, "parents": [folder_id]},
        media_body=media,
        fields="id, webViewLink",
    ).execute()
 
    print(f"[INFO] Uploaded to Google Drive: {file.get('webViewLink')}")
    return file.get("webViewLink", "")
def load_ledger() -> dict:
    if LEDGER_PATH.exists():
        with open(LEDGER_PATH) as f:
            return json.load(f)
    return {"processed": {}}
 
def save_ledger(ledger: dict):
    with open(LEDGER_PATH, "w") as f:
        json.dump(ledger, f, indent=2)
 
def mark_processed(video_id: str, metadata: dict):
    ledger = load_ledger()
    ledger["processed"][video_id] = {
        "title": metadata.get("title", ""),
        "processed_at": datetime.now(timezone.utc).isoformat(),
        "drive_link": metadata.get("drive_link", ""),
    }
    save_ledger(ledger)
 
 
# ---------------------------------------------------------------------------
 
# ---------------------------------------------------------------------------
# DOCX Rendering — Oakleaf branded
# ---------------------------------------------------------------------------
 
OL_GREEN_HEX = "185430"
OL_GREEN_LIGHT_HEX = "397D52"
OL_GRAY_TEXT_HEX = "404040"
OL_GRAY_LIGHT_HEX = "F2F2F2"
 
 
def _parse_summary_sections(summary_text: str) -> list[dict]:
    """
    Split Claude's summary into named sections.
    Returns list of {heading, body} dicts.
    """
    section_markers = [
        "EXECUTIVE SUMMARY",
        "KEY VOTES & ACTIONS",
        "TOPIC FLAGS",
        "UPCOMING & FOLLOW-UP",
    ]
    sections = []
    remaining = summary_text.strip()
 
    for i, marker in enumerate(section_markers):
        pattern = re.compile(
            r"(?:^|\n)\s*(?:\d+\.\s*)?" + re.escape(marker) + r"\s*\n",
            re.IGNORECASE,
        )
        m = pattern.search(remaining)
        if not m:
            continue
        before = remaining[: m.start()].strip()
        if before and sections:
            sections[-1]["body"] += "\n" + before
        elif before:
            sections.append({"heading": "", "body": before})
        next_start = len(remaining)
        for next_marker in section_markers[i + 1 :]:
            np = re.compile(
                r"(?:^|\n)\s*(?:\d+\.\s*)?" + re.escape(next_marker) + r"\s*\n",
                re.IGNORECASE,
            )
            nm = np.search(remaining, m.end())
            if nm and nm.start() < next_start:
                next_start = nm.start()
        body = remaining[m.end() : next_start].strip()
        sections.append({"heading": marker, "body": body})
        remaining = remaining[next_start:]
 
    if remaining.strip() and sections:
        sections[-1]["body"] += "\n" + remaining.strip()
 
    return sections
 
 
def render_docx(doc: dict) -> bytes:
    """
    Render the summary dict as a branded Oakleaf .docx file.
    Returns raw docx bytes.
    """
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
 
    document = DocxDocument()
 
    # -- Page setup --
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
 
    green = RGBColor(0x18, 0x54, 0x30)
    gray_text = RGBColor(0x40, 0x40, 0x40)
 
    # -- Header --
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = hp.add_run("Oakleaf reDevelopment")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = green
    hp.add_run("  |  ").font.size = Pt(9)
    title_run = hp.add_run(doc["title"])
    title_run.font.size = Pt(9)
    title_run.font.color.rgb = gray_text
 
    # -- Footer --
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("Portland City Council Monitor  |  oakleaf.dev  |  Confidential — Internal Use Only")
    fr.font.size = Pt(7)
    fr.font.color.rgb = gray_text
 
    # -- Title --
    title_para = document.add_heading(doc["title"], level=1)
    for run in title_para.runs:
        run.font.color.rgb = green
 
    # -- Meta table --
    meta_rows = [
        ("Date", doc["upload_date"]),
        ("Duration", f"{doc['duration_min']} minutes"),
        ("Source", doc["youtube_url"]),
        ("Transcription", doc["transcription_method"]),
        ("Processed", doc["processed_at"]),
    ]
    table = document.add_table(rows=len(meta_rows), cols=2)
    table.style = "Light Grid Accent 1"
    for idx, (label, value) in enumerate(meta_rows):
        row = table.rows[idx]
        cell_label = row.cells[0]
        cell_value = row.cells[1]
        cell_label.text = label
        cell_value.text = value
        for paragraph in cell_label.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = green
        for paragraph in cell_value.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = gray_text
 
    document.add_paragraph()  # spacer
 
    # -- Body sections --
    sections = _parse_summary_sections(doc["summary"])
 
    if not sections:
        for para in doc["summary"].split("\n\n"):
            para = para.strip()
            if para:
                p = document.add_paragraph()
                _add_formatted_text(p, para, gray_text)
    else:
        for section in sections:
            if section["heading"]:
                h = document.add_heading(section["heading"], level=2)
                for run in h.runs:
                    run.font.color.rgb = green
 
            for line in section["body"].split("\n"):
                line = line.strip()
                if not line:
                    continue
                if line.startswith(("- ", "* ", "\u2022 ")):
                    text = line[2:].strip()
                    p = document.add_paragraph(style="List Bullet")
                    _add_formatted_text(p, text, gray_text)
                elif re.match(r"^\d+\.\s", line):
                    text = re.sub(r"^\d+\.\s*", "", line)
                    p = document.add_paragraph(style="List Bullet")
                    _add_formatted_text(p, text, gray_text)
                else:
                    p = document.add_paragraph()
                    _add_formatted_text(p, line, gray_text)
 
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()
 
 
def _add_formatted_text(paragraph, text: str, default_color):
    """Add text to a paragraph, converting **bold** markers to actual bold runs."""
    from docx.shared import Pt
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = default_color
        else:
            run = paragraph.add_run(part)
            run.font.size = Pt(10)
            run.font.color.rgb = default_color
 
 
def build_markdown_backup(doc: dict) -> str:
    """Plain markdown version saved to repo as a readable backup."""
    return (
        f"# Portland City Council Meeting Summary\n"
        f"**Date:** {doc['upload_date']}\n"
        f"**Title:** {doc['title']}\n"
        f"**Duration:** {doc['duration_min']} minutes\n"
        f"**Source:** {doc['youtube_url']}\n"
        f"**Transcription:** {doc['transcription_method']}\n"
        f"**Processed:** {doc['processed_at']}\n\n"
        f"---\n\n"
        f"{doc['summary']}\n"
    )
 
 
# Build summary document
# ---------------------------------------------------------------------------
 
def build_summary_doc(
    summary: str,
    metadata: dict,
    video_id: str,
    transcription_method: str,
) -> dict:
    """Return a structured dict used by both the PDF renderer and the repo markdown backup."""
    youtube_url = f"https://www.youtube.com/watch?v={video_id}"
    return {
        "summary": summary,
        "title": metadata.get("title", "Portland City Council Meeting"),
        "upload_date": metadata.get("upload_date", "Unknown"),
        "duration_min": metadata.get("duration", 0) // 60,
        "youtube_url": youtube_url,
        "transcription_method": transcription_method,
        "processed_at": datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC"),
    }
 
 
def save_and_upload(doc: dict, video_id: str, date_str: str, metadata: dict) -> str:
    """Render DOCX, save locally, commit markdown backup to repo, upload DOCX to Drive."""
    summary_dir = Path("summaries")
    summary_dir.mkdir(exist_ok=True)
 
    title = metadata.get("title", video_id)
    safe_title = re.sub(r'[<>:"/\\|?*]', '', title).strip()
    base_name = f"{date_str}_{safe_title}_summary"
 
    # --- Render DOCX ---
    docx_bytes = render_docx(doc)
    docx_path = summary_dir / f"{base_name}.docx"
    with open(docx_path, "wb") as f:
        f.write(docx_bytes)
    print(f"  DOCX saved: {docx_path}")
 
    # --- Markdown backup for repo (keeps git history readable) ---
    md_content = build_markdown_backup(doc)
    md_path = summary_dir / f"{base_name}.md"
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(md_content)
    commit_summary_to_repo([docx_path, md_path], metadata)
 
    # --- Upload DOCX to Google Drive ---
    drive_link = ""
    drive_folder_id = os.environ.get("GOOGLE_DRIVE_FOLDER_ID")
    if drive_folder_id:
        try:
            service = get_drive_service()
            if service:
                drive_link = upload_to_drive(
                    service, f"{base_name}.docx", docx_bytes, drive_folder_id
                )
        except Exception as e:
            print(f"[WARN] Drive upload failed: {e}")
 
    mark_processed(video_id, {"title": metadata["title"], "drive_link": drive_link})
    return drive_link
def process_from_dispatch():
    config = load_config()
 
    payload_json = os.environ.get("TRANSCRIPT_PAYLOAD", "{}")
    payload = json.loads(payload_json)
 
    video_id = payload.get("video_id")
    if not video_id:
        print("[ERROR] No video_id in dispatch payload.")
        sys.exit(1)
 
    # Read transcript and metadata from repo files (committed by local script)
    transcript_path = payload.get("transcript_path", f"transcripts/{video_id}/transcript.txt")
    metadata_path = payload.get("metadata_path", f"transcripts/{video_id}/metadata.json")
 
    repo_root = SCRIPT_DIR
    transcript_file = repo_root / transcript_path
    metadata_file = repo_root / metadata_path
 
    if not transcript_file.exists():
        print(f"[ERROR] Transcript file not found: {transcript_file}")
        print("  Make sure the local script committed it to the repo.")
        sys.exit(1)
 
    with open(transcript_file, encoding="utf-8") as f:
        transcript = f.read()
 
    if not transcript:
        print("[ERROR] Transcript file is empty.")
        sys.exit(1)
 
    # Load metadata from committed file, with fallbacks to dispatch payload
    if metadata_file.exists():
        with open(metadata_file, encoding="utf-8") as f:
            meta_from_file = json.load(f)
    else:
        print(f"[WARN] Metadata file not found: {metadata_file}, using dispatch payload.")
        meta_from_file = {}
 
    metadata = {
        "id": video_id,
        "title": meta_from_file.get("title", payload.get("title", "Unknown")),
        "upload_date": meta_from_file.get("upload_date", payload.get("upload_date", "Unknown")),
        "duration": meta_from_file.get("duration", payload.get("duration", 0)),
    }
 
    # Parse chapters from metadata file
    chapters_text = ""
    chapters = meta_from_file.get("chapters", [])
    if chapters:
        ch_lines = [
            f"  [{format_timestamp(ch.get('start', 0))}] {ch.get('headline', '')}: {ch.get('summary', '')}"
            for ch in chapters
        ]
        chapters_text = "\n".join(ch_lines)
 
    assemblyai_summary = meta_from_file.get("assemblyai_summary", "")
 
    print(f"\n=== Processing transcript from repo ===")
    print(f"  Video: {metadata['title']}")
    print(f"  Date:  {metadata['upload_date']}")
    print(f"  Transcript: {len(transcript)} chars")
 
    prompt = build_prompt(transcript, metadata, config["watch_topics"], chapters_text, assemblyai_summary)
    summary = summarize_with_claude(prompt)
    doc = build_summary_doc(summary, metadata, video_id, "AssemblyAI (speaker diarization enabled)")
    drive_link = save_and_upload(doc, video_id, metadata["upload_date"], metadata)
 
    print(f"\n=== DONE ===")
    if drive_link:
        print(f"  Drive: {drive_link}")
 
 
# ---------------------------------------------------------------------------
# Process from URL (YouTube captions fallback)
# ---------------------------------------------------------------------------
 
def process_from_url(youtube_url: str):
    config = load_config()
    video_id = extract_video_id(youtube_url)
 
    print("\n=== STEP 1: Fetching video metadata ===")
    metadata = get_video_metadata(video_id)
    print(f"  Title: {metadata['title']}")
    print(f"  Date:  {metadata['upload_date']}")
 
    print("\n=== STEP 2: Fetching YouTube captions ===")
    transcript = fetch_youtube_captions(video_id)
    print(f"  Transcript: {len(transcript)} chars")
 
    print("\n=== STEP 3: Summarizing with Claude ===")
    prompt = build_prompt(transcript, metadata, config["watch_topics"])
    summary = summarize_with_claude(prompt)
    doc = build_summary_doc(summary, metadata, video_id, "YouTube auto-captions (no speaker diarization)")
    drive_link = save_and_upload(doc, video_id, metadata["upload_date"], metadata)
 
    print(f"\n=== DONE ===")
    if drive_link:
        print(f"  Drive: {drive_link}")
 
 
# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------
 
def main():
    parser = argparse.ArgumentParser(description="Council Meeting Summarizer")
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument("--from-dispatch", action="store_true")
    group.add_argument("--url", help="Process a YouTube URL (uses captions)")
 
    args = parser.parse_args()
 
    if args.from_dispatch:
        process_from_dispatch()
    elif args.url:
        process_from_url(args.url)
 
 
if __name__ == "__main__":
    main()
 
